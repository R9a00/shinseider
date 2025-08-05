from fastapi.testclient import TestClient
from docx import Document
import os
import sys
sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))

from main import app

client = TestClient(app)

def test_generate_business_plan_word_document():
    test_data = {
        "business_summary": "これは事業概要のテストです。",
        "market_analysis": "市場分析のテストです。",
        "competitive_advantage": "競合優位性のテストです。",
        "products_services": "製品・サービスのテストです。",
        "marketing_strategy": "マーケティング戦略のテストです。",
        "revenue_plan": "収益計画のテストです。",
        "funding_plan": "資金計画のテストです。",
        "implementation_structure": "実施体制のテストです。",
    }
    response = client.post("/generate_business_plan", json=test_data)

    assert response.status_code == 200
    assert response.headers["content-type"] == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    import re
    assert re.match(r"attachment; filename\*?=utf-8''%E4%BA%8B%E6%A5%AD%E8%A8%88%E7%94%BB%E6%9B%B8\.docx", response.headers["content-disposition"])

    # ダウンロードされたWordファイルを一時的に保存し、内容を検証
    with open("test_business_plan.docx", "wb") as f:
        f.write(response.content)

    document = Document("test_business_plan.docx")

    # 要約セクションの確認
    summary_found = False
    for paragraph in document.paragraphs:
        if "【事業概要】" in paragraph.text and "【市場分析】" in paragraph.text:
            summary_found = True
            break
    assert summary_found, "Summary section not found or incomplete in the generated Word document."

    # 各セクションのタイトルと内容の確認
    expected_headings = [
        "事業概要", "市場分析", "競合優位性", "製品・サービス",
        "マーケティング戦略", "収益計画", "資金計画", "実施体制"
    ]
    for heading in expected_headings:
        found_heading = False
        for paragraph in document.paragraphs:
            if paragraph.text == heading:
                found_heading = True
                break
        assert found_heading, f"Heading '{heading}' not found in the generated Word document."

    # 生成された一時ファイルを削除
    os.remove("test_business_plan.docx")
