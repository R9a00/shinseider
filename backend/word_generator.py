from docx import Document
from docx.shared import Inches
import json
from docx.enum.text import WD_ALIGN_PARAGRAPH

def generate_word(file_path: str, content: str):
    document = Document()

    # contentはJSON文字列として渡されることを想定
    try:
        parsed_content = json.loads(content)
    except json.JSONDecodeError:
        document.add_paragraph(content)
        document.save(file_path)
        return

    # ファイル名に基づいて異なる処理を行う
    if "textbook" in file_path: # Phase1の診断レポート
        document.add_heading('事業承継診断レポート', 0)

        sections = {
            "はじめに": "このレポートは、アトツギ甲子園出場に向けた事業構想を整理するための診断結果をまとめたものです。",
            "事業アイデアの核と既存事業とのシナジー": "",
            "社会課題解決と提供価値": "",
            "後継者の想い・ストーリー": "",
            "経営資源の活用戦略": "",
            "ターゲット市場と成長性": "",
            "競合優位性と差別化戦略": "",
            "課題とリスク、そして克服戦略": "",
            "事業実現に向けた目標とロードマップ": "",
            "アトツギ甲子園への目標と事業承継への展望": "",
        }

        # 質問と回答をセクションにマッピング
        if len(parsed_content) > 0: sections["事業アイデアの核と既存事業とのシナジー"] = parsed_content[0].get('answer', '')
        if len(parsed_content) > 1: sections["社会課題解決と提供価値"] = parsed_content[1].get('answer', '')
        if len(parsed_content) > 2: sections["後継者の想い・ストーリー"] = parsed_content[2].get('answer', '')
        if len(parsed_content) > 3: sections["経営資源の活用戦略"] = parsed_content[3].get('answer', '')
        if len(parsed_content) > 4: sections["ターゲット市場と成長性"] = parsed_content[4].get('answer', '')
        if len(parsed_content) > 5: sections["競合優位性と差別化戦略"] = parsed_content[5].get('answer', '')
        if len(parsed_content) > 6: sections["課題とリスク、そして克服戦略"] = parsed_content[6].get('answer', '')
        if len(parsed_content) > 7: sections["事業実現に向けた目標とロードマップ"] = parsed_content[7].get('answer', '')
        if len(parsed_content) > 8: sections["アトツギ甲子園への目標と事業承継への展望"] = parsed_content[8].get('answer', '')

        # ドキュメントにセクションを追加
        for title, text in sections.items():
            document.add_heading(title, level=1)
            if text:
                for paragraph_text in text.split('\n'):
                    document.add_paragraph(paragraph_text)
            else:
                document.add_paragraph("該当する回答がありませんでした。")
            document.add_paragraph() # セクション間の区切り

    elif "business_plan" in file_path: # 事業計画書
        document.add_heading('事業計画書', 0)

        sections_map = {
            "business_summary": "事業概要",
            "market_analysis": "市場分析",
            "competitive_advantage": "競合優位性",
            "products_services": "製品・サービス",
            "marketing_strategy": "マーケティング戦略",
            "revenue_plan": "収益計画",
            "funding_plan": "資金計画",
            "implementation_structure": "実施体制",
        }

        for key, title in sections_map.items():
            document.add_heading(title, level=1)
            text_content = parsed_content.get(key, "").strip()
            if text_content:
                for paragraph_text in text_content.split('\n'):
                    document.add_paragraph(paragraph_text)
            else:
                document.add_paragraph("該当する情報がありませんでした。")
            document.add_paragraph() # セクション間の区切り

    document.save(file_path)


