import sys
import os

# プロジェクトルートをパスに追加
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from fastapi import FastAPI, Request, HTTPException, Depends, File, UploadFile, Form
from fastapi.responses import FileResponse, JSONResponse
# from word_generator import generate_word
import os
import json
import yaml
from typing import Dict, Any
from fastapi.middleware.cors import CORSMiddleware
from middleware.security import SecurityMiddleware, ErrorHandlingMiddleware, RateLimitMiddleware
from models import DesireRequest, ApplicationAdviceRequest, TextbookRequest, BusinessPlanRequest, ContactRequest
from typing import Optional
from secure_file_utils import get_secure_file_manager
from config import settings
import logging

BASE_DIR = os.path.dirname(os.path.abspath(__file__))

# セキュリティログの設定
logging.basicConfig(level=getattr(logging, settings.LOG_LEVEL))
security_logger = logging.getLogger("security")

def load_subsidy_data():
    with open(settings.SUBSIDIES_CONFIG_PATH, "r", encoding="utf-8") as f:
        return yaml.safe_load(f)

def load_version_history():
    version_history_path = os.path.join(BASE_DIR, "version_history.yaml")
    try:
        with open(version_history_path, "r", encoding="utf-8") as f:
            return yaml.safe_load(f)
    except FileNotFoundError:
        return {"metadata": {}, "subsidies": {}}

app = FastAPI()

origins = settings.CORS_ORIGINS + [
    "http://localhost:3000",  # Development fallback
    "http://127.0.0.1:3333",
    "http://127.0.0.1:3000"
]

# セキュリティミドルウェアの追加
app.add_middleware(
    CORSMiddleware,
    allow_origins=origins,
    allow_credentials=True,
    allow_methods=["GET", "POST", "OPTIONS"],
    allow_headers=["*"],
)

# セキュリティミドルウェア（REQ-SEC-001, REQ-SEC-006, REQ-SEC-010）
app.add_middleware(SecurityMiddleware, max_request_size=10*1024*1024)  # 10MB制限（ファイル添付のため増加）

# エラーハンドリングミドルウェア（REQ-SEC-004）
is_production = os.getenv("ENVIRONMENT") == "production"
app.add_middleware(ErrorHandlingMiddleware, is_production=is_production)

# レート制限ミドルウェア（REQ-SEC-008）
app.add_middleware(RateLimitMiddleware, max_requests=100, window_seconds=60)

@app.get("/")
def read_root():
    return {"Hello": "World"}

@app.get("/subsidies")
async def get_subsidies():
    """利用可能な補助金の一覧を取得します。"""
    subsidies_data = load_subsidy_data()
    subsidies_list = [
        {"id": subsidy.get("id"), "name": subsidy.get("name")}
        for subsidy in subsidies_data
    ]
    return {"subsidies": subsidies_list}

@app.get("/subsidies/{subsidy_id}/metadata")
async def get_subsidy_metadata(subsidy_id: str):
    """指定された補助金のメタデータを取得します。"""
    subsidies_data = load_subsidy_data()
    subsidy = next((s for s in subsidies_data if s.get("id") == subsidy_id), None)
    if not subsidy:
        raise HTTPException(status_code=404, detail="Subsidy not found")
    return {"id": subsidy.get("id"), "name": subsidy.get("name")}

@app.get("/version-history")
async def get_version_history():
    """全体のバージョン履歴情報を取得します。"""
    try:
        version_data = load_version_history()
        return version_data
    except Exception as e:
        security_logger.error(f"Failed to load version history: {str(e)}")
        raise HTTPException(status_code=500, detail="Failed to load version history")

@app.get("/subsidies/{subsidy_id}/version-history")
async def get_subsidy_version_history(subsidy_id: str):
    """指定された補助金のバージョン履歴を取得します。"""
    try:
        version_data = load_version_history()
        subsidies_data = load_subsidy_data()
        
        # 補助金が存在するかチェック
        subsidy = next((s for s in subsidies_data if s.get("id") == subsidy_id), None)
        if not subsidy:
            raise HTTPException(status_code=404, detail="Subsidy not found")
        
        # バージョン履歴を取得
        subsidy_version = version_data.get("subsidies", {}).get(subsidy_id, {})
        if not subsidy_version:
            return {
                "subsidy_id": subsidy_id,
                "subsidy_name": subsidy.get("name"),
                "version": "未設定",
                "last_updated": "未設定",
                "source_references": [],
                "change_history": []
            }
        
        return {
            "subsidy_id": subsidy_id,
            "subsidy_name": subsidy.get("name"),
            **subsidy_version
        }
        
    except HTTPException:
        raise
    except Exception as e:
        security_logger.error(f"Failed to load version history for {subsidy_id}: {str(e)}")
        raise HTTPException(status_code=500, detail="Failed to load version history")

@app.get("/get_application_questions/{subsidy_id}")
async def get_application_questions(subsidy_id: str):
    """指定された補助金の申請書作成に必要な質問項目（セクション）を取得します。"""
    subsidies_data = load_subsidy_data()
    subsidy = next((s for s in subsidies_data if s.get("id") == subsidy_id), None)
    if not subsidy:
        raise HTTPException(status_code=404, detail="Subsidy not found")

    sections = subsidy.get("sections", [])
    validation = subsidy.get("validation", {})
    checklist = subsidy.get("checklist", [])
    tasks = subsidy.get("tasks", {})
    
    return {
        "sections": sections,
        "validation": validation,
        "checklist": checklist,
        "tasks": tasks
    }

@app.post("/generate_application_advice")
async def generate_application_advice(advice_request: ApplicationAdviceRequest):
    """ユーザーの回答と目的に応じて、多様なアウトプットを生成します。"""
    try:
        subsidies_data = load_subsidy_data()
        subsidy = next((s for s in subsidies_data if s.get("id") == advice_request.subsidy_id), None)
        if not subsidy:
            raise HTTPException(status_code=404, detail="Subsidy not found")

        prompt_template = subsidy.get("llm_prompt_template")
        if not prompt_template:
            raise HTTPException(status_code=500, detail="Prompt template for this subsidy is not configured.")

        # 「わからない」状態を質問文に変換する関数
        def convert_unknown_to_question(value, task_label, section_title):
            if isinstance(value, str) and "わからない・要相談" in value:
                return f"【質問】{task_label or section_title}について、何を書けばよいか具体的に教えてください。書き方のポイントや例があれば教えてください。"
            elif isinstance(value, list):
                converted = []
                for item in value:
                    if isinstance(item, str) and "わからない・要相談" in item:
                        converted.append(f"【質問】{task_label or section_title}について、どのような内容を入力すればよいか具体例とともに教えてください。")
                    elif isinstance(item, dict):
                        converted_item = {}
                        for k, v in item.items():
                            if isinstance(v, str) and "わからない・要相談" in v:
                                converted_item[k] = f"【質問】{task_label or section_title}の{k}について、どのような内容を設定すればよいか、具体的な例を交えて教えてください。"
                            else:
                                converted_item[k] = v
                        converted.append(converted_item)
                    else:
                        converted.append(item)
                return converted
            return value

        # ユーザーの回答を整形（出場資格情報を除外）
        def format_answers(answers, sections, exclude_eligibility=False):
            formatted = ""
            # 出場資格関連のタスクID（除外対象）
            eligibility_tasks = {"MINI_001_AGE", "MINI_002_UNDER_39", "MINI_005_NAME", "MINI_006_COMPANY", "MINI_003_SUCCESSION_PLAN"}
            
            for section in sections:
                section_id = section.get("id")
                if section_id in answers:
                    title = section.get("title", "不明な項目")
                    
                    # ミニタスク形式のデータ構造をチェック
                    section_data = answers.get(section_id)
                    if isinstance(section_data, dict):
                        # 出場資格関連のタスクを除外するかチェック
                        filtered_tasks = []
                        for task_id, task_value in section_data.items():
                            if exclude_eligibility and task_id in eligibility_tasks:
                                continue  # 出場資格情報をスキップ
                            if task_value:
                                # 「わからない・要相談」を質問文に変換
                                task_label = task_id
                                if section.get("input_modes", {}).get("micro_tasks"):
                                    for task in section["input_modes"]["micro_tasks"]:
                                        if task.get("task_id") == task_id:
                                            task_label = task.get("label", task_id)
                                            break
                                converted_value = convert_unknown_to_question(task_value, task_label, title)
                                filtered_tasks.append((task_id, converted_value))
                        
                        if filtered_tasks:  # 有効なタスクがある場合のみセクションを追加
                            formatted += f"▼ {title}\n"
                            
                            # ミニタスクの場合（各タスクIDがキーになっている）
                            for task_id, task_value in filtered_tasks:
                                # セクション内のタスクを見つける
                                task_label = task_id
                                if section.get("input_modes", {}).get("micro_tasks"):
                                    for task in section["input_modes"]["micro_tasks"]:
                                        if task.get("task_id") == task_id:
                                            task_label = task.get("label", task_id)
                                            break
                                
                                if isinstance(task_value, list):
                                    # 配列の場合
                                    if task_value:  # 空でない場合
                                        formatted += f"  {task_label}: {', '.join(str(v) for v in task_value if v)}\n"
                                else:
                                    formatted += f"  {task_label}: {task_value}\n"
                            formatted += "\n"
                    else:
                        # 通常の文字列データの場合
                        if section_data:
                            # 「わからない・要相談」を質問文に変換
                            converted_data = convert_unknown_to_question(section_data, None, title)
                            formatted += f"▼ {title}\n"
                            formatted += f"A: {converted_data}\n\n"
            return formatted

        # targetに応じて異なる出力を生成
        if advice_request.target == "ai":
            # AI相談用：全情報を含めたプロンプトテンプレートを使用
            formatted_answers = format_answers(advice_request.answers, subsidy.get("sections", []), exclude_eligibility=False)
            final_prompt = prompt_template.format(
                user_answers=formatted_answers,
                input_mode=advice_request.input_mode
            )
            security_logger.info(f"LLM prompt generated for subsidy: {advice_request.subsidy_id}")
            return {"output": final_prompt, "type": "prompt"}
            
        elif advice_request.target == "human":
            # 専門家相談用：出場資格情報を除外して具体的な質問を生成
            formatted_answers = format_answers(advice_request.answers, subsidy.get("sections", []), exclude_eligibility=True)
            
            expert_consultation = f"""専門家への相談事項

事業内容
{formatted_answers}

お聞きしたいこと

1. この事業アイデアの独自性をもっと強くアピールするには、どのような点を強調すべきでしょうか？

2. ターゲット市場の規模や成長性を示すために、どのようなデータを調査・収集すべきでしょうか？

3. この事業を実現するために見落としている準備や課題があれば教えてください。

4. 収益予測について、業界水準や市場実態と比較してのご意見をお聞かせください。

5. 初期投資や運転資金の調達で、効果的な方法があれば教えてください。

6. 申請書類で最も効果的にアピールする構成や表現方法を教えてください。

7. この補助金の審査で特に重視されるポイントがあれば教えてください。"""
            
            security_logger.info(f"Human consultation summary generated for subsidy: {advice_request.subsidy_id}")
            return {"output": expert_consultation, "type": "summary"}
            
        elif advice_request.target == "self":
            # 自己評価用：出場資格情報を除外して問いかけを生成
            formatted_answers = format_answers(advice_request.answers, subsidy.get("sections", []), exclude_eligibility=True)
            
            self_reflection = f"""申請書の自己チェックリスト

あなたの入力内容
{formatted_answers}

自分で確認してみましょう

事業アイデアの魅力度チェック
□ 似たようなサービス・事業と比べて、明確な違いを3つ説明できますか？
□ ターゲット顧客が「これは欲しい」と思う理由を具体的に説明できますか？
□ 「どれくらいの人が使ってくれそうか」を数字で説明できますか？

実現可能性チェック
□ この事業を成功させるのに必要なスキルや経験は身についていますか？
□ 人・物・お金の準備はできていますか？不足分の調達方法は決まっていますか？
□ マイルストーンの日程は現実的で達成可能ですか？

収益計画チェック
□ 売上予測の根拠は説得力がありますか？「なんとなく」になっていませんか？
□ 必要な費用をもれなく計算できていますか？
□ きちんと利益が出る計画になっていますか？

伝わりやすさチェック
□ 事業内容を30秒で説明できますか？
□ なぜあなたがこの事業をやるのか、心に響くストーリーがありますか？
□ 主張を裏付ける具体的な数字やデータが入っていますか？

改善のヒント
上記でチェックが付かない項目について、以下を試してみてください：

1. 足りない情報やデータを調べる
2. 感覚的な表現を具体的な数字に変える  
3. 家族や友人に説明してみて反応を確認
4. 似たサービスをもう一度詳しく調べる
5. 上記で不安な点があれば専門家に相談する

最終確認
全ての項目にチェックが付いたら、あなたの申請書は審査員にしっかり伝わる内容になっています。"""
            
            security_logger.info(f"Self-reflection questions generated for subsidy: {advice_request.subsidy_id}")
            return {"output": self_reflection, "type": "reflection"}
        else:
            raise HTTPException(status_code=400, detail="Invalid target specified.")

    except Exception as e:
        security_logger.error(f"Failed to generate application advice: {e}")
        raise HTTPException(status_code=500, detail="Failed to generate application advice")

@app.post("/save_desire")
async def save_desire(desire_request: DesireRequest):
    """事業承継者の回答をセキュアに保存（REQ-SEC-002, REQ-SEC-006）"""
    try:
        secure_file_manager = get_secure_file_manager()
        
        formatted_desire = ""
        for item in desire_request.answers:
            formatted_desire += f"Q: {item.question}\nA: {item.answer}\n\n"

        # セキュアなファイル書き込み
        secure_file_manager.secure_file_write(formatted_desire, "desire.txt")
        
        security_logger.info("Desire saved securely")
        return {"message": "Desire saved successfully"}
    except Exception as e:
        security_logger.error(f"Failed to save desire: {e}")
        raise HTTPException(status_code=500, detail="Failed to save desire")

@app.post("/generate_textbook")
async def generate_textbook(textbook_request: TextbookRequest):
    """教科書をセキュアに生成（REQ-SEC-002, REQ-SEC-006）"""
    try:
        from docx import Document
        from docx.shared import Inches
        import tempfile
        import json
        
        security_logger.info("Textbook generation requested")
        
        # 新しいWord文書を作成
        doc = Document()
        
        # タイトルを追加
        title = doc.add_heading(getattr(textbook_request, 'title', '事業承継診断レポート'), 0)
        
        # 生成日時を追加
        doc.add_paragraph(f"生成日時: {str(__import__('datetime').datetime.now().strftime('%Y年%m月%d日 %H:%M:%S'))}")
        doc.add_paragraph("")
        
        # コンテンツをパース
        try:
            content_data = json.loads(textbook_request.content)
            if isinstance(content_data, list):
                for item in content_data:
                    if isinstance(item, dict) and 'question' in item and 'answer' in item:
                        # 質問をヘディングとして追加
                        doc.add_heading(f"Q: {item['question']}", level=2)
                        # 回答を段落として追加
                        doc.add_paragraph(f"A: {item['answer'] or '（回答なし）'}")
                        doc.add_paragraph("")
            else:
                doc.add_paragraph(str(content_data))
        except (json.JSONDecodeError, TypeError):
            # JSONでない場合はそのまま追加
            doc.add_paragraph(textbook_request.content)
        
        # 一時ファイルに保存
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        doc.save(temp_file.name)
        temp_file.close()
        
        return FileResponse(
            temp_file.name,
            media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            filename="事業承継診断レポート.docx"
        )
        
    except Exception as e:
        security_logger.error(f"Failed to generate textbook: {e}")
        raise HTTPException(status_code=500, detail="Failed to generate textbook")

@app.post("/save_application_data")
async def save_application_data(data: Dict[str, Any]):
    """申請データをWord文書形式で保存・ダウンロード"""
    try:
        from docx import Document
        from docx.shared import Inches
        import tempfile
        from datetime import datetime
        
        security_logger.info("Application data save requested")
        
        # 新しいWord文書を作成
        doc = Document()
        
        # タイトル
        title = data.get('subsidy_name', '補助金') + ' 申請準備書'
        doc.add_heading(title, 0)
        
        # 保存情報
        doc.add_paragraph(f"保存日時: {datetime.now().strftime('%Y年%m月%d日 %H:%M:%S')}")
        doc.add_paragraph(f"進捗状況: {data.get('progress', 0)}%完了")
        doc.add_paragraph("")
        
        # 回答データを整理して追加
        answers = data.get("answers", {})
        for section_id, section_data in answers.items():
            if isinstance(section_data, dict):
                # ミニタスク形式の場合
                section_title = section_id.replace("_", " ").title()
                doc.add_heading(f"■ {section_title}", level=1)
                
                for task_id, task_value in section_data.items():
                    if task_value:
                        # タスクラベルを整形
                        task_label = task_id.replace("MINI_", "").replace("_", " ").title()
                        
                        if isinstance(task_value, list):
                            if task_value:  # 空でない場合
                                # マイルストーンの特別処理
                                if task_id == "MILESTONES":
                                    doc.add_paragraph(f"◆ 事業マイルストーン")
                                    
                                    # 有効なマイルストーンのみ抽出
                                    valid_milestones = [item for item in task_value 
                                                      if isinstance(item, dict) and any(v for v in item.values() if v)]
                                    
                                    if valid_milestones:
                                        # 表を作成
                                        table = doc.add_table(rows=1, cols=3)
                                        table.style = 'Table Grid'
                                        
                                        # ヘッダー行
                                        hdr_cells = table.rows[0].cells
                                        hdr_cells[0].text = '時期'
                                        hdr_cells[1].text = '達成内容'
                                        hdr_cells[2].text = '備考'
                                        
                                        # ヘッダー行を太字に
                                        for cell in hdr_cells:
                                            for paragraph in cell.paragraphs:
                                                for run in paragraph.runs:
                                                    run.bold = True
                                        
                                        # データ行を追加
                                        for item in valid_milestones:
                                            row_cells = table.add_row().cells
                                            row_cells[0].text = item.get('ym', '')
                                            row_cells[1].text = item.get('note', '')
                                            row_cells[2].text = item.get('owner', '')
                                        
                                        doc.add_paragraph("")  # 表の後にスペース
                                else:
                                    doc.add_paragraph(f"◆ {task_label}:")
                                    for item in task_value:
                                        if isinstance(item, dict):
                                            # 他の構造化データの場合
                                            for key, value in item.items():
                                                if value:
                                                    doc.add_paragraph(f"  - {key}: {value}", style='List Bullet')
                                        else:
                                            doc.add_paragraph(f"  - {item}", style='List Bullet')
                        else:
                            doc.add_paragraph(f"◆ {task_label}: {task_value}")
                
                doc.add_paragraph("")  # セクション間のスペース
                
            elif section_data:
                # 通常の文字列データの場合
                section_title = section_id.replace("_", " ").title()
                doc.add_heading(f"■ {section_title}", level=1)
                doc.add_paragraph(section_data)
                doc.add_paragraph("")
        
        # フッター情報
        doc.add_paragraph("")
        doc.add_paragraph("=" * 50)
        doc.add_paragraph("このデータは「シンセイダー」で作成されました。")
        doc.add_paragraph("プライバシー保護: このデータはサーバーに保存されていません。")
        
        # 一時ファイルに保存
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        doc.save(temp_file.name)
        temp_file.close()
        
        filename = f"申請準備書_{data.get('subsidy_name', '補助金')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        
        return FileResponse(
            temp_file.name,
            media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            filename=filename
        )
        
    except Exception as e:
        security_logger.error(f"Failed to save application data: {e}")
        raise HTTPException(status_code=500, detail="Failed to save application data")

@app.post("/generate_business_plan")
async def generate_business_plan(business_plan_request: BusinessPlanRequest):
    """事業計画書をセキュアに生成（REQ-SEC-002, REQ-SEC-006）"""
    try:
        from docx import Document
        import tempfile
        
        security_logger.info("Business plan generation requested")
        
        # 新しいWord文書を作成
        doc = Document()
        
        # タイトル
        doc.add_heading('事業計画書', 0)
        
        # 生成日時
        doc.add_paragraph(f"生成日時: {str(__import__('datetime').datetime.now().strftime('%Y年%m月%d日 %H:%M:%S'))}")
        doc.add_paragraph("")
        
        # 各セクションを追加
        sections = [
            ('事業概要', business_plan_request.business_summary),
            ('市場分析', business_plan_request.market_analysis),
            ('競合優位性', business_plan_request.competitive_advantage),
            ('製品・サービス', business_plan_request.products_services),
            ('マーケティング戦略', business_plan_request.marketing_strategy),
            ('収益計画', business_plan_request.revenue_plan),
            ('資金計画', business_plan_request.funding_plan),
            ('実施体制', business_plan_request.implementation_structure)
        ]
        
        for title, content in sections:
            if content:
                doc.add_heading(title, level=1)
                doc.add_paragraph(content)
                doc.add_paragraph("")
        
        # 一時ファイルに保存
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        doc.save(temp_file.name)
        temp_file.close()
        
        return FileResponse(
            temp_file.name,
            media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            filename="事業計画書.docx"
        )
        
    except Exception as e:
        security_logger.error(f"Failed to generate business plan: {e}")
        raise HTTPException(status_code=500, detail="Failed to generate business plan")

@app.post("/send_contact")
async def send_contact(
    name: str = Form(...),
    email: str = Form(...),
    subject: str = Form(...),
    message: str = Form(...),
    attachment: Optional[UploadFile] = File(None)
):
    """お問い合わせメールを送信（ファイル添付対応）"""
    try:
        import aiosmtplib
        from email.message import EmailMessage
        from email.mime.multipart import MIMEMultipart
        from email.mime.text import MIMEText
        from email.mime.base import MIMEBase
        from email import encoders
        import os
        import tempfile
        
        security_logger.info("Contact form submission received")
        
        # 入力値検証
        dangerous_chars = ['<', '>', '"', "'", '&', 'script', 'javascript']
        for field_value in [name, subject, message]:
            for char in dangerous_chars:
                if char.lower() in field_value.lower():
                    raise HTTPException(status_code=400, detail="不正な文字が含まれています")
        
        # Gmail SMTP設定
        SMTP_HOST = "smtp.gmail.com"
        SMTP_PORT = 587
        SMTP_USER = os.getenv("GMAIL_USER")
        SMTP_PASS = os.getenv("GMAIL_APP_PASSWORD")
        TO_EMAIL = SMTP_USER
        
        if not SMTP_USER or not SMTP_PASS:
            raise HTTPException(status_code=500, detail="Mail configuration not found")
        
        # メール作成
        msg = MIMEMultipart()
        msg["From"] = SMTP_USER
        msg["To"] = TO_EMAIL
        msg["Subject"] = f"【シンセイダー】{subject}"
        
        # メール本文
        body_text = f"""シンセイダーからのお問い合わせです。

お名前: {name}
メールアドレス: {email}
件名: {subject}

メッセージ:
{message}

---
送信日時: {__import__('datetime').datetime.now().strftime('%Y年%m月%d日 %H:%M:%S')}
"""
        msg.attach(MIMEText(body_text, 'plain', 'utf-8'))
        
        # ファイル添付処理
        if attachment and attachment.filename:
            # セキュリティチェック
            allowed_extensions = {'.pdf', '.doc', '.docx', '.xls', '.xlsx', '.jpg', '.jpeg', '.png', '.gif'}
            file_ext = os.path.splitext(attachment.filename.lower())[1]
            
            if file_ext not in allowed_extensions:
                raise HTTPException(status_code=400, detail="許可されていないファイル形式です")
            
            if attachment.size and attachment.size > 5 * 1024 * 1024:  # 5MB制限
                raise HTTPException(status_code=400, detail="ファイルサイズが大きすぎます（5MB以下にしてください）")
            
            # ファイル内容読み込み
            file_content = await attachment.read()
            
            # 基本的なセキュリティチェック（マジックナンバー確認）
            magic_numbers = {
                b'\x25\x50\x44\x46': '.pdf',
                b'\xd0\xcf\x11\xe0': '.doc/.xls',
                b'\x50\x4b\x03\x04': '.docx/.xlsx',
                b'\xff\xd8\xff': '.jpg',
                b'\x89\x50\x4e\x47': '.png',
                b'\x47\x49\x46': '.gif'
            }
            
            is_valid = any(file_content.startswith(magic) for magic in magic_numbers.keys())
            if not is_valid:
                raise HTTPException(status_code=400, detail="ファイルの形式が正しくありません")
            
            # 添付ファイル追加
            part = MIMEBase('application', 'octet-stream')
            part.set_payload(file_content)
            encoders.encode_base64(part)
            
            # ファイル名の適切なエンコーディング
            from email.utils import formataddr
            import urllib.parse
            
            # ファイル名をRFC2047形式でエンコード
            encoded_filename = urllib.parse.quote(attachment.filename.encode('utf-8'))
            part.add_header(
                'Content-Disposition',
                f'attachment; filename*=UTF-8\'\'{encoded_filename}'
            )
            msg.attach(part)
            
            security_logger.info(f"File attached: {attachment.filename} ({len(file_content)} bytes)")
        
        # メール送信
        await aiosmtplib.send(
            msg,
            hostname=SMTP_HOST,
            port=SMTP_PORT,
            start_tls=True,
            username=SMTP_USER,
            password=SMTP_PASS
        )
        
        security_logger.info(f"Contact email sent successfully from {email}")
        return {"message": "お問い合わせを送信しました", "success": True}
        
    except HTTPException:
        raise
    except Exception as e:
        security_logger.error(f"Failed to send contact email: {e}")
        raise HTTPException(status_code=500, detail="メール送信に失敗しました")

@app.get("/news")
async def get_news():
    """ニュース・お知らせを取得"""
    try:
        news_path = os.path.join(BASE_DIR, "news_content.yaml")
        with open(news_path, 'r', encoding='utf-8') as file:
            news_data = yaml.safe_load(file)
        return news_data
    except FileNotFoundError:
        raise HTTPException(status_code=404, detail="ニュースファイルが見つかりません")
    except Exception as e:
        security_logger.error(f"Failed to load news: {str(e)}")
        raise HTTPException(status_code=500, detail=f"ニュースデータの取得に失敗しました: {str(e)}")

@app.get("/knowledge-base")
async def get_knowledge_base():
    """補助金基礎知識を取得"""
    try:
        knowledge_path = os.path.join(BASE_DIR, "knowledge_base.yaml")
        with open(knowledge_path, 'r', encoding='utf-8') as file:
            knowledge_data = yaml.safe_load(file)
        return knowledge_data
    except FileNotFoundError:
        raise HTTPException(status_code=404, detail="基礎知識ファイルが見つかりません")
    except Exception as e:
        security_logger.error(f"Failed to load knowledge base: {str(e)}")
        raise HTTPException(status_code=500, detail=f"基礎知識データの取得に失敗しました: {str(e)}")

@app.get("/subsidy-selection")
async def get_subsidy_selection():
    """補助金選択画面用のデータを取得"""
    try:
        subsidies_data = load_subsidy_data()
        # 補助金一覧を整理して返す
        subsidies_list = []
        for subsidy in subsidies_data:
            subsidies_list.append({
                "id": subsidy.get("id"),
                "name": subsidy.get("name"),
                "category": subsidy.get("category", "一般"),
                "description": subsidy.get("description", ""),
                "max_amount": subsidy.get("max_amount", ""),
                "subsidy_rate": subsidy.get("subsidy_rate", ""),
                "difficulty": subsidy.get("difficulty", "普通"),
                "suitable_for": subsidy.get("suitable_for", [])
            })
        return {"subsidies": subsidies_list}
    except Exception as e:
        security_logger.error(f"Failed to load subsidy selection: {str(e)}")
        raise HTTPException(status_code=500, detail=f"補助金選択データの取得に失敗しました: {str(e)}")

@app.get("/knowledge-base/{section_id}")
async def get_knowledge_section(section_id: str):
    """特定の基礎知識セクションを取得"""
    try:
        knowledge_path = os.path.join(BASE_DIR, "knowledge_base.yaml")
        with open(knowledge_path, 'r', encoding='utf-8') as file:
            knowledge_data = yaml.safe_load(file)
        
        if section_id not in knowledge_data.get('content', {}):
            raise HTTPException(status_code=404, detail=f"セクション '{section_id}' が見つかりません")
        
        return {
            'section_id': section_id,
            'data': knowledge_data['content'][section_id],
            'metadata': knowledge_data.get('metadata', {})
        }
    except HTTPException:
        raise
    except Exception as e:
        security_logger.error(f"Failed to load knowledge section: {str(e)}")
        raise HTTPException(status_code=500, detail=f"基礎知識データの取得に失敗しました: {str(e)}")

if __name__ == "__main__":
    import uvicorn
    import os
    port = int(os.environ.get("PORT", 8888))
    uvicorn.run(app, host="0.0.0.0", port=port)